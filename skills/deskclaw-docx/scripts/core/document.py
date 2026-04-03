"""Document management: create, read, copy, merge, convert to PDF."""

import json
import shutil
from pathlib import Path

try:
    import docx
    from docx import Document
except ImportError:
    docx = None
    Document = None


def _check_docx():
    if Document is None:
        raise RuntimeError("python-docx is required. Install with: pip install python-docx")


def create_document(filename, title=None, author=None):
    _check_docx()
    doc = Document()
    core = doc.core_properties
    if title is not None:
        core.title = str(title)
    if author is not None:
        core.author = str(author)
    doc.save(filename)
    return f"Created document: {filename}"


def get_document_info(filename):
    _check_docx()
    doc = Document(filename)
    core = doc.core_properties
    paragraphs = list(doc.paragraphs)
    tables = list(doc.tables)
    info = {
        "filename": filename,
        "title": getattr(core, "title", None) or "",
        "author": getattr(core, "author", None) or "",
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
    }
    return json.dumps(info, ensure_ascii=False, indent=2)


def get_document_text(filename):
    _check_docx()
    doc = Document(filename)
    parts = []
    for p in doc.paragraphs:
        parts.append(p.text)
    return "\n".join(parts) if parts else ""


def get_document_outline(filename):
    _check_docx()
    doc = Document(filename)
    outline = []
    for i, p in enumerate(doc.paragraphs):
        style = p.style.name if p.style else ""
        if style.startswith("Heading") or style == "Title":
            level = 0
            if style == "Title":
                level = 0
            elif "Heading" in style:
                try:
                    level = int(style.replace("Heading", "").strip() or "1")
                except ValueError:
                    level = 1
            outline.append({"index": i, "level": level, "text": p.text[:80]})
    return json.dumps(outline, ensure_ascii=False, indent=2)


def list_available_documents(directory="."):
    path = Path(directory).resolve()
    if not path.is_dir():
        return json.dumps({"error": f"Not a directory: {directory}"})
    files = sorted(path.glob("*.docx"))
    return json.dumps([str(f) for f in files], ensure_ascii=False, indent=2)


def copy_document(source_filename, destination_filename=None):
    src = Path(source_filename)
    if not src.exists():
        return f"Error: file not found: {source_filename}"
    if destination_filename is None:
        destination_filename = str(src.parent / f"{src.stem}_copy{src.suffix}")
    shutil.copy2(source_filename, destination_filename)
    return f"Copied to: {destination_filename}"


def convert_to_pdf(filename, output_filename=None):
    try:
        from docx2pdf import convert
    except ImportError:
        return "Error: docx2pdf is required for PDF conversion. Install with: pip install docx2pdf"
    src = Path(filename)
    if not src.exists():
        return f"Error: file not found: {filename}"
    if output_filename is None:
        output_filename = str(src.with_suffix(".pdf"))
    convert(str(src), output_filename)
    return f"Converted to: {output_filename}"


def merge_documents(filename, source_files):
    _check_docx()
    if not source_files:
        return "Error: source_files cannot be empty"
    doc = Document(source_files[0])
    for path in source_files[1:]:
        sub = Document(path)
        for element in sub.element.body:
            doc.element.body.append(element)
    doc.save(filename)
    return f"Merged {len(source_files)} documents into: {filename}"


def add_header(filename, text, font_name=None, font_size=None, bold=None, alignment="center"):
    """Add header text to the document."""
    _check_docx()
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document(filename)
    section = doc.sections[0]
    header = section.header
    
    if header.paragraphs:
        p = header.paragraphs[0]
    else:
        p = header.add_paragraph()
    
    p.clear()
    run = p.add_run(text)
    
    if font_name:
        run.font.name = font_name
    if font_size:
        run.font.size = Pt(font_size / 2.0)
    if bold is not None:
        run.font.bold = bold
    
    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    p.alignment = align_map.get(alignment, WD_ALIGN_PARAGRAPH.CENTER)
    
    doc.save(filename)
    return f"Added header: {text[:30]}"


def add_footer(filename, text, font_name=None, font_size=None, bold=None, alignment="center"):
    """Add footer text to the document."""
    _check_docx()
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document(filename)
    section = doc.sections[0]
    footer = section.footer
    
    if footer.paragraphs:
        p = footer.paragraphs[0]
    else:
        p = footer.add_paragraph()
    
    p.clear()
    run = p.add_run(text)
    
    if font_name:
        run.font.name = font_name
    if font_size:
        run.font.size = Pt(font_size / 2.0)
    if bold is not None:
        run.font.bold = bold
    
    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    p.alignment = align_map.get(alignment, WD_ALIGN_PARAGRAPH.CENTER)
    
    doc.save(filename)
    return f"Added footer: {text[:30]}"


def add_page_number(filename, position="footer", alignment="center", format_text="第 {page} 页"):
    """Add page numbers to header or footer.
    
    format_text can include {page} for current page number.
    """
    _check_docx()
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    doc = Document(filename)
    section = doc.sections[0]
    
    if position == "header":
        target = section.header
    else:
        target = section.footer
    
    if target.paragraphs:
        p = target.paragraphs[0]
        if p.text.strip():
            p = target.add_paragraph()
    else:
        p = target.add_paragraph()
    
    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    p.alignment = align_map.get(alignment, WD_ALIGN_PARAGRAPH.CENTER)
    
    # Split format_text by {page} and add field
    parts = format_text.split("{page}")
    if parts[0]:
        p.add_run(parts[0])
    
    # Add PAGE field
    run = p.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char_begin)
    
    instr_text = OxmlElement("w:instrText")
    instr_text.text = "PAGE"
    run._r.append(instr_text)
    
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_end)
    
    if len(parts) > 1 and parts[1]:
        p.add_run(parts[1])
    
    doc.save(filename)
    return f"Added page number to {position}"


def set_different_first_page_header(filename, enabled=True):
    """Enable or disable different first page header/footer."""
    _check_docx()
    doc = Document(filename)
    section = doc.sections[0]
    section.different_first_page_header_footer = enabled
    doc.save(filename)
    return f"Different first page header/footer: {'enabled' if enabled else 'disabled'}"


def add_header_image(filename, image_path, width=None, alignment="center"):
    """Add an image to the header."""
    _check_docx()
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document(filename)
    section = doc.sections[0]
    header = section.header
    
    if header.paragraphs:
        p = header.paragraphs[0]
    else:
        p = header.add_paragraph()
    
    run = p.add_run()
    if width:
        run.add_picture(image_path, width=Inches(width))
    else:
        run.add_picture(image_path)
    
    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    p.alignment = align_map.get(alignment, WD_ALIGN_PARAGRAPH.CENTER)
    
    doc.save(filename)
    return f"Added image to header: {image_path}"


def add_table_of_contents(filename, title="目录", max_level=3, position="start"):
    """Add a table of contents to the document.
    
    title: TOC title (default "目录")
    max_level: maximum heading level to include (1-9)
    position: "start" to insert at beginning, or paragraph index
    
    Note: The TOC field needs to be updated in Word/WPS to show actual content.
    """
    _check_docx()
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document(filename)
    body = doc.element.body
    
    # Create title paragraph element if needed
    title_p = None
    if title:
        title_para = doc.add_paragraph(title, style='Heading 1')
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p = title_para._p
    
    # Create TOC paragraph
    toc_para = doc.add_paragraph()
    toc_p = toc_para._p
    
    # Create TOC field
    run = toc_para.add_run()
    
    # Begin field
    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fld_char_begin)
    
    # Field instruction
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = f' TOC \\o "1-{max_level}" \\h \\z \\u '
    run._r.append(instr_text)
    
    # Separate field
    fld_char_separate = OxmlElement('w:fldChar')
    fld_char_separate.set(qn('w:fldCharType'), 'separate')
    run._r.append(fld_char_separate)
    
    # Placeholder text
    placeholder_run = toc_para.add_run("请在 Word 中右键点击此处，选择'更新域'以生成目录")
    
    # End field
    end_run = toc_para.add_run()
    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')
    end_run._r.append(fld_char_end)
    
    # Create page break paragraph
    page_break_para = doc.add_paragraph()
    page_break_para.add_run().add_break(docx.enum.text.WD_BREAK.PAGE)
    page_break_p = page_break_para._p
    
    # Now move elements to the correct position
    # First remove them from their current position (end of document)
    if title_p is not None:
        body.remove(title_p)
    body.remove(toc_p)
    body.remove(page_break_p)
    
    # Insert at the specified position
    if position == "start":
        # Insert in order: title, toc, page_break
        body.insert(0, page_break_p)
        body.insert(0, toc_p)
        if title_p is not None:
            body.insert(0, title_p)
    elif isinstance(position, int):
        para_elements = [e for e in body if e.tag.endswith('}p')]
        if position < len(para_elements):
            insert_point = list(body).index(para_elements[position])
            # Insert in reverse order at same point
            body.insert(insert_point, page_break_p)
            body.insert(insert_point, toc_p)
            if title_p is not None:
                body.insert(insert_point, title_p)
        else:
            # Position out of range, append at end
            if title_p is not None:
                body.append(title_p)
            body.append(toc_p)
            body.append(page_break_p)
    else:
        # Default: append at end
        if title_p is not None:
            body.append(title_p)
        body.append(toc_p)
        body.append(page_break_p)
    
    doc.save(filename)
    return f"Added table of contents (levels 1-{max_level})"


def update_table_of_contents(filename):
    """Mark TOC for update. 
    
    Note: Actual update happens when the document is opened in Word/WPS.
    This function sets the dirty flag on the TOC field.
    """
    _check_docx()
    from docx.oxml.ns import qn
    
    doc = Document(filename)
    
    # Find and mark TOC fields as dirty
    for para in doc.paragraphs:
        for run in para.runs:
            for elem in run._r:
                if elem.tag.endswith('}fldChar'):
                    fld_type = elem.get(qn('w:fldCharType'))
                    if fld_type == 'begin':
                        elem.set(qn('w:dirty'), 'true')
    
    doc.save(filename)
    return "Marked TOC for update (will refresh when opened in Word/WPS)"


def set_page_size(filename, width=None, height=None, paper="A4"):
    """Set page size.
    
    paper: 'A4', 'A3', 'Letter', 'Legal', or 'custom'
    width/height: in inches (only used when paper='custom')
    """
    _check_docx()
    from docx.shared import Inches, Mm
    
    # Predefined paper sizes (in mm), keyed by uppercase for case-insensitive lookup
    paper_sizes = {
        "A4": (210, 297),
        "A3": (297, 420),
        "LETTER": (215.9, 279.4),
        "LEGAL": (215.9, 355.6),
        "B5": (176, 250),
    }
    
    doc = Document(filename)
    
    for section in doc.sections:
        if paper.upper() in paper_sizes:
            w_mm, h_mm = paper_sizes[paper.upper()]
            section.page_width = Mm(w_mm)
            section.page_height = Mm(h_mm)
        elif paper.lower() == "custom" and width and height:
            section.page_width = Inches(width)
            section.page_height = Inches(height)
        else:
            return f"Error: Unknown paper size '{paper}'. Use A4, A3, Letter, Legal, B5, or custom."
    
    doc.save(filename)
    return f"Set page size to {paper}"


def set_page_margins(filename, top=None, bottom=None, left=None, right=None):
    """Set page margins in inches."""
    _check_docx()
    from docx.shared import Inches
    
    doc = Document(filename)
    
    for section in doc.sections:
        if top is not None:
            section.top_margin = Inches(top)
        if bottom is not None:
            section.bottom_margin = Inches(bottom)
        if left is not None:
            section.left_margin = Inches(left)
        if right is not None:
            section.right_margin = Inches(right)
    
    doc.save(filename)
    return f"Set page margins"


def set_page_orientation(filename, orientation="portrait"):
    """Set page orientation.
    
    orientation: 'portrait' (纵向) or 'landscape' (横向)
    """
    _check_docx()
    from docx.enum.section import WD_ORIENT
    
    doc = Document(filename)
    
    for section in doc.sections:
        current_width = section.page_width
        current_height = section.page_height
        
        if orientation.lower() == "landscape":
            section.orientation = WD_ORIENT.LANDSCAPE
            # Swap width and height if needed
            if current_width < current_height:
                section.page_width = current_height
                section.page_height = current_width
        else:  # portrait
            section.orientation = WD_ORIENT.PORTRAIT
            # Swap width and height if needed
            if current_width > current_height:
                section.page_width = current_height
                section.page_height = current_width
    
    doc.save(filename)
    return f"Set page orientation to {orientation}"


def get_page_settings(filename):
    """Get current page settings."""
    _check_docx()
    import json
    from docx.shared import Inches
    
    doc = Document(filename)
    section = doc.sections[0]
    
    # Convert EMUs to inches for readability
    def emu_to_inches(emu):
        if emu is None:
            return None
        return round(emu / 914400, 2)
    
    settings = {
        "page_width_inches": emu_to_inches(section.page_width),
        "page_height_inches": emu_to_inches(section.page_height),
        "orientation": "landscape" if section.orientation == 1 else "portrait",
        "margins": {
            "top": emu_to_inches(section.top_margin),
            "bottom": emu_to_inches(section.bottom_margin),
            "left": emu_to_inches(section.left_margin),
            "right": emu_to_inches(section.right_margin),
        }
    }
    
    return json.dumps(settings, ensure_ascii=False, indent=2)


def add_watermark(filename, text, font_name="Arial", font_size=48, color="C0C0C0", diagonal=True):
    """Add a text watermark to the document.
    
    text: watermark text (e.g., "机密", "DRAFT", "CONFIDENTIAL")
    font_size: in points
    color: hex color (light gray recommended)
    diagonal: if True, text is rotated diagonally
    """
    _check_docx()
    from docx.oxml.ns import qn, nsmap
    from docx.oxml import OxmlElement
    from docx.shared import Pt, RGBColor
    
    doc = Document(filename)
    
    # Watermark needs to be added to the header
    for section in doc.sections:
        header = section.header
        
        # Create a paragraph for the watermark
        if not header.paragraphs:
            p = header.add_paragraph()
        else:
            p = header.paragraphs[0]
        
        # Create VML shape for watermark (Word uses VML for watermarks)
        # We'll use a simpler approach: add semi-transparent text
        
        # Create the watermark using a shape in the header
        # This is a simplified version - true watermarks use VML
        
        # For now, we'll add it as a styled text in header
        # A proper implementation would use VML shapes
        
        run = p.add_run(text)
        run.font.size = Pt(font_size)
        run.font.name = font_name
        
        # Set color
        from docx.shared import RGBColor
        c = color.lstrip('#')
        if len(c) == 6:
            r, g, b = int(c[0:2], 16), int(c[2:4], 16), int(c[4:6], 16)
            run.font.color.rgb = RGBColor(r, g, b)
        
        # Center the watermark
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.save(filename)
    return f"Added watermark: {text}"


def add_watermark_image(filename, image_path, width=None):
    """Add an image watermark to the document header."""
    _check_docx()
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document(filename)
    
    for section in doc.sections:
        header = section.header
        
        if not header.paragraphs:
            p = header.add_paragraph()
        else:
            p = header.paragraphs[0]
        
        run = p.add_run()
        if width:
            run.add_picture(image_path, width=Inches(width))
        else:
            run.add_picture(image_path)
        
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.save(filename)
    return f"Added image watermark: {image_path}"


def add_footnote(filename, paragraph_index, footnote_text):
    """Add a footnote to a paragraph.
    
    Note: python-docx has limited footnote support. This creates a basic footnote.
    """
    _check_docx()
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import zipfile
    import tempfile
    import shutil
    import os
    from pathlib import Path
    import xml.etree.ElementTree as ET
    
    path = Path(filename)
    
    # Get existing footnotes count
    footnote_id = "1"
    with zipfile.ZipFile(path, 'r') as z:
        if "word/footnotes.xml" in z.namelist():
            with z.open("word/footnotes.xml") as f:
                tree = ET.parse(f)
                root = tree.getroot()
                ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                footnotes = root.findall(".//w:footnote", ns)
                # Find max id (skip -1 and 0 which are separator and continuation)
                max_id = 0
                for fn in footnotes:
                    fn_id = fn.get(f"{{{ns['w']}}}id", "0")
                    try:
                        if int(fn_id) > max_id:
                            max_id = int(fn_id)
                    except ValueError:
                        pass
                footnote_id = str(max_id + 1)
    
    with tempfile.TemporaryDirectory() as tmpdir:
        # Extract docx
        with zipfile.ZipFile(path, 'r') as z:
            z.extractall(tmpdir)
        
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        ET.register_namespace('w', ns['w'])
        
        # Check if footnotes.xml exists, create if not
        footnotes_path = os.path.join(tmpdir, "word", "footnotes.xml")
        if not os.path.exists(footnotes_path):
            # Create new footnotes.xml with required separator and continuation footnotes
            footnotes_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
    <w:footnote w:type="separator" w:id="-1">
        <w:p><w:r><w:separator/></w:r></w:p>
    </w:footnote>
    <w:footnote w:type="continuationSeparator" w:id="0">
        <w:p><w:r><w:continuationSeparator/></w:r></w:p>
    </w:footnote>
</w:footnotes>'''
            with open(footnotes_path, 'w', encoding='utf-8') as f:
                f.write(footnotes_xml)
            
            # Update [Content_Types].xml
            content_types_path = os.path.join(tmpdir, "[Content_Types].xml")
            ct_tree = ET.parse(content_types_path)
            ct_root = ct_tree.getroot()
            ct_ns = "http://schemas.openxmlformats.org/package/2006/content-types"
            
            has_footnotes = False
            for override in ct_root.findall(f".//{{{ct_ns}}}Override"):
                if override.get("PartName") == "/word/footnotes.xml":
                    has_footnotes = True
                    break
            
            if not has_footnotes:
                override = ET.SubElement(ct_root, f"{{{ct_ns}}}Override")
                override.set("PartName", "/word/footnotes.xml")
                override.set("ContentType", "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml")
                ct_tree.write(content_types_path, xml_declaration=True, encoding='UTF-8')
            
            # Update document.xml.rels
            rels_path = os.path.join(tmpdir, "word", "_rels", "document.xml.rels")
            rels_tree = ET.parse(rels_path)
            rels_root = rels_tree.getroot()
            rels_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
            
            # Find next rId
            max_rid = 0
            for rel in rels_root.findall(f".//{{{rels_ns}}}Relationship"):
                rid = rel.get("Id", "rId0")
                try:
                    rid_num = int(rid.replace("rId", ""))
                    max_rid = max(max_rid, rid_num)
                except ValueError:
                    pass
            
            new_rid = f"rId{max_rid + 1}"
            rel = ET.SubElement(rels_root, f"{{{rels_ns}}}Relationship")
            rel.set("Id", new_rid)
            rel.set("Type", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes")
            rel.set("Target", "footnotes.xml")
            rels_tree.write(rels_path, xml_declaration=True, encoding='UTF-8')
        
        # Add footnote to footnotes.xml
        fn_tree = ET.parse(footnotes_path)
        fn_root = fn_tree.getroot()
        
        footnote = ET.SubElement(fn_root, f"{{{ns['w']}}}footnote")
        footnote.set(f"{{{ns['w']}}}id", footnote_id)
        
        p = ET.SubElement(footnote, f"{{{ns['w']}}}p")
        r = ET.SubElement(p, f"{{{ns['w']}}}r")
        t = ET.SubElement(r, f"{{{ns['w']}}}t")
        t.text = footnote_text
        
        fn_tree.write(footnotes_path, xml_declaration=True, encoding='UTF-8')
        
        # Add footnote reference to document.xml
        doc_path = os.path.join(tmpdir, "word", "document.xml")
        doc_tree = ET.parse(doc_path)
        doc_root = doc_tree.getroot()
        
        paragraphs = doc_root.findall(f".//{{{ns['w']}}}p")
        if paragraph_index >= len(paragraphs):
            return f"Error: paragraph_index {paragraph_index} out of range"
        
        target_p = paragraphs[paragraph_index]
        
        # Add footnote reference at the end of the paragraph
        r = ET.SubElement(target_p, f"{{{ns['w']}}}r")
        footnote_ref = ET.SubElement(r, f"{{{ns['w']}}}footnoteReference")
        footnote_ref.set(f"{{{ns['w']}}}id", footnote_id)
        
        doc_tree.write(doc_path, xml_declaration=True, encoding='UTF-8')
        
        # Repack
        output_path = str(path) + ".tmp"
        with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as z:
            for root_dir, dirs, files in os.walk(tmpdir):
                for file in files:
                    file_path = os.path.join(root_dir, file)
                    arcname = os.path.relpath(file_path, tmpdir)
                    z.write(file_path, arcname)
        
        shutil.move(output_path, path)
    
    return f"Added footnote to paragraph {paragraph_index}"


def add_endnote(filename, paragraph_index, endnote_text):
    """Add an endnote to a paragraph.
    
    Similar to footnote but appears at the end of the document.
    """
    _check_docx()
    from pathlib import Path
    import zipfile
    import tempfile
    import shutil
    import os
    import xml.etree.ElementTree as ET
    
    path = Path(filename)
    
    # Get existing endnotes count
    endnote_id = "1"
    with zipfile.ZipFile(path, 'r') as z:
        if "word/endnotes.xml" in z.namelist():
            with z.open("word/endnotes.xml") as f:
                tree = ET.parse(f)
                root = tree.getroot()
                ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                endnotes = root.findall(".//w:endnote", ns)
                max_id = 0
                for en in endnotes:
                    en_id = en.get(f"{{{ns['w']}}}id", "0")
                    try:
                        if int(en_id) > max_id:
                            max_id = int(en_id)
                    except ValueError:
                        pass
                endnote_id = str(max_id + 1)
    
    with tempfile.TemporaryDirectory() as tmpdir:
        with zipfile.ZipFile(path, 'r') as z:
            z.extractall(tmpdir)
        
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        ET.register_namespace('w', ns['w'])
        
        endnotes_path = os.path.join(tmpdir, "word", "endnotes.xml")
        if not os.path.exists(endnotes_path):
            endnotes_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:endnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
    <w:endnote w:type="separator" w:id="-1">
        <w:p><w:r><w:separator/></w:r></w:p>
    </w:endnote>
    <w:endnote w:type="continuationSeparator" w:id="0">
        <w:p><w:r><w:continuationSeparator/></w:r></w:p>
    </w:endnote>
</w:endnotes>'''
            with open(endnotes_path, 'w', encoding='utf-8') as f:
                f.write(endnotes_xml)
            
            content_types_path = os.path.join(tmpdir, "[Content_Types].xml")
            ct_tree = ET.parse(content_types_path)
            ct_root = ct_tree.getroot()
            ct_ns = "http://schemas.openxmlformats.org/package/2006/content-types"
            
            has_endnotes = False
            for override in ct_root.findall(f".//{{{ct_ns}}}Override"):
                if override.get("PartName") == "/word/endnotes.xml":
                    has_endnotes = True
                    break
            
            if not has_endnotes:
                override = ET.SubElement(ct_root, f"{{{ct_ns}}}Override")
                override.set("PartName", "/word/endnotes.xml")
                override.set("ContentType", "application/vnd.openxmlformats-officedocument.wordprocessingml.endnotes+xml")
                ct_tree.write(content_types_path, xml_declaration=True, encoding='UTF-8')
            
            rels_path = os.path.join(tmpdir, "word", "_rels", "document.xml.rels")
            rels_tree = ET.parse(rels_path)
            rels_root = rels_tree.getroot()
            rels_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
            
            max_rid = 0
            for rel in rels_root.findall(f".//{{{rels_ns}}}Relationship"):
                rid = rel.get("Id", "rId0")
                try:
                    rid_num = int(rid.replace("rId", ""))
                    max_rid = max(max_rid, rid_num)
                except ValueError:
                    pass
            
            new_rid = f"rId{max_rid + 1}"
            rel = ET.SubElement(rels_root, f"{{{rels_ns}}}Relationship")
            rel.set("Id", new_rid)
            rel.set("Type", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes")
            rel.set("Target", "endnotes.xml")
            rels_tree.write(rels_path, xml_declaration=True, encoding='UTF-8')
        
        en_tree = ET.parse(endnotes_path)
        en_root = en_tree.getroot()
        
        endnote = ET.SubElement(en_root, f"{{{ns['w']}}}endnote")
        endnote.set(f"{{{ns['w']}}}id", endnote_id)
        
        p = ET.SubElement(endnote, f"{{{ns['w']}}}p")
        r = ET.SubElement(p, f"{{{ns['w']}}}r")
        t = ET.SubElement(r, f"{{{ns['w']}}}t")
        t.text = endnote_text
        
        en_tree.write(endnotes_path, xml_declaration=True, encoding='UTF-8')
        
        doc_path = os.path.join(tmpdir, "word", "document.xml")
        doc_tree = ET.parse(doc_path)
        doc_root = doc_tree.getroot()
        
        paragraphs = doc_root.findall(f".//{{{ns['w']}}}p")
        if paragraph_index >= len(paragraphs):
            return f"Error: paragraph_index {paragraph_index} out of range"
        
        target_p = paragraphs[paragraph_index]
        
        r = ET.SubElement(target_p, f"{{{ns['w']}}}r")
        endnote_ref = ET.SubElement(r, f"{{{ns['w']}}}endnoteReference")
        endnote_ref.set(f"{{{ns['w']}}}id", endnote_id)
        
        doc_tree.write(doc_path, xml_declaration=True, encoding='UTF-8')
        
        output_path = str(path) + ".tmp"
        with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as z:
            for root_dir, dirs, files in os.walk(tmpdir):
                for file in files:
                    file_path = os.path.join(root_dir, file)
                    arcname = os.path.relpath(file_path, tmpdir)
                    z.write(file_path, arcname)
        
        shutil.move(output_path, path)
    
    return f"Added endnote to paragraph {paragraph_index}"
