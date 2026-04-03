"""Content addition: headings, paragraphs, tables, pictures, lists."""

import json
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    Document = None


def _check_docx():
    if Document is None:
        raise RuntimeError("python-docx is required. Install with: pip install python-docx")


def _hex_to_rgb(hex_color):
    h = hex_color.lstrip("#")
    if len(h) == 6:
        return tuple(int(h[i : i + 2], 16) for i in (0, 2, 4))
    return (0, 0, 0)


def _apply_run_font(run, font_name=None, font_size=None, bold=None, italic=None, color=None):
    if font_name is not None:
        run.font.name = font_name
    if font_size is not None:
        run.font.size = Pt(font_size / 2.0)
    if bold is not None:
        run.font.bold = bool(bold)
    if italic is not None:
        run.font.italic = bool(italic)
    if color is not None:
        c = color.lstrip("#")
        run.font.color.rgb = None
        from docx.shared import RGBColor
        try:
            run.font.color.rgb = RGBColor(*_hex_to_rgb(c))
        except Exception:
            pass


def add_heading(filename, text, level=1, font_name=None, font_size=None, bold=None, italic=None, border_bottom=False):
    _check_docx()
    doc = Document(filename)
    p = doc.add_heading(text, level=min(max(level, 0), 9))
    if font_name is not None or font_size is not None or bold is not None or italic is not None:
        for run in p.runs:
            _apply_run_font(run, font_name, font_size, bold, italic, None)
    if border_bottom:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        pBdr.append(bottom)
        pPr.append(pBdr)
    doc.save(filename)
    return f"Added heading (level {level}): {text[:50]}"


def add_paragraph(filename, text, style=None, font_name=None, font_size=None, bold=None, italic=None, color=None):
    _check_docx()
    doc = Document(filename)
    p = doc.add_paragraph(text, style=style)
    for run in p.runs:
        _apply_run_font(run, font_name, font_size, bold, italic, color)
    doc.save(filename)
    return f"Added paragraph: {text[:50]}..."


def add_table(filename, rows, cols, data=None):
    _check_docx()
    doc = Document(filename)
    table = doc.add_table(rows=rows, cols=cols)
    if data:
        for r, row_data in enumerate(data):
            if r >= rows:
                break
            row = table.rows[r]
            for c, cell_text in enumerate(row_data):
                if c < len(row.cells):
                    row.cells[c].text = str(cell_text)
    doc.save(filename)
    return f"Added table: {rows}x{cols}"


def add_picture(filename, image_path, width=None):
    _check_docx()
    doc = Document(filename)
    if width is not None:
        from docx.shared import Inches
        doc.add_picture(image_path, width=Inches(width))
    else:
        doc.add_picture(image_path)
    doc.save(filename)
    return f"Added picture: {image_path}"


def add_page_break(filename):
    _check_docx()
    doc = Document(filename)
    doc.add_page_break()
    doc.save(filename)
    return "Added page break"


def _find_paragraph_index_by_text(doc, target_text):
    for i, p in enumerate(doc.paragraphs):
        if target_text in p.text:
            return i
    return None


def insert_header_near_text(filename, target_text=None, header_title=None, position="after", header_style="Heading 1", target_paragraph_index=None):
    _check_docx()
    doc = Document(filename)
    if target_paragraph_index is not None:
        idx = target_paragraph_index
    elif target_text:
        idx = _find_paragraph_index_by_text(doc, target_text)
        if idx is None:
            return f"Error: text not found: {target_text}"
    else:
        return "Error: provide target_text or target_paragraph_index"
    if position == "before":
        insert_idx = idx
    else:
        insert_idx = idx + 1
    new_p = doc.add_heading(header_title or "Heading", level=1)
    if header_style and header_style != "Heading 1":
        new_p.style = header_style
    body = doc.element.body
    target_el = body[idx].__copy__()
    body.insert(insert_idx, body[insert_idx - 1].__copy__())
    for i, p in enumerate(doc.paragraphs):
        if i == len(doc.paragraphs) - 1 and insert_idx <= len(body):
            break
    doc.save(filename)
    return f"Inserted header '{header_title}' {position} target"


def insert_line_or_paragraph_near_text(filename, target_text=None, line_text=None, position="after", line_style=None, target_paragraph_index=None):
    _check_docx()
    doc = Document(filename)
    if target_paragraph_index is not None:
        idx = target_paragraph_index
    elif target_text:
        idx = _find_paragraph_index_by_text(doc, target_text)
        if idx is None:
            return f"Error: text not found: {target_text}"
    else:
        return "Error: provide target_text or target_paragraph_index"
    new_p = doc.add_paragraph(line_text or "", style=line_style)
    body = doc.element.body
    para_elements = [e for e in body if e.tag.endswith("}p")]
    if idx + (1 if position == "after" else 0) <= len(para_elements):
        body.insert(body.index(para_elements[idx]) + (1 if position == "after" else 0), new_p._p)
    doc.save(filename)
    return f"Inserted paragraph {position} target"


def insert_numbered_list_near_text(filename, target_text=None, list_items=None, position="after", target_paragraph_index=None, bullet_type="bullet"):
    _check_docx()
    if not list_items:
        return "Error: list_items required"
    doc = Document(filename)
    if target_paragraph_index is not None:
        idx = target_paragraph_index
    elif target_text:
        idx = _find_paragraph_index_by_text(doc, target_text)
        if idx is None:
            return f"Error: text not found: {target_text}"
    else:
        idx = 0
    style = "List Bullet" if bullet_type == "bullet" else "List Number"
    for i, item in enumerate(list_items):
        p = doc.add_paragraph(item, style=style)
    doc.save(filename)
    return f"Inserted {len(list_items)} list items"


def _get_unique_cells(table):
    """Get unique cells in a table, handling merged cells.
    
    Returns a dict: {(row, col): cell} where each cell appears only once.
    Also returns a mapping of all (row, col) to their "master" (row, col).
    """
    unique_cells = {}
    cell_to_master = {}
    seen_cells = set()
    
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            cell_id = id(cell)
            if cell_id not in seen_cells:
                seen_cells.add(cell_id)
                unique_cells[(ri, ci)] = cell
                cell_to_master[(ri, ci)] = (ri, ci)
            else:
                # Find the master cell
                for (mr, mc), c in unique_cells.items():
                    if id(c) == cell_id:
                        cell_to_master[(ri, ci)] = (mr, mc)
                        break
    
    return unique_cells, cell_to_master


def get_table_info(filename, table_index=None, show_merged=True):
    """Get information about tables in the document."""
    _check_docx()
    doc = Document(filename)
    tables = doc.tables
    
    if table_index is not None:
        if table_index >= len(tables):
            return json.dumps({"error": f"Table index {table_index} out of range (total: {len(tables)})"})
        table = tables[table_index]
        
        unique_cells, cell_to_master = _get_unique_cells(table)
        
        rows_info = []
        for ri, row in enumerate(table.rows):
            cells_info = []
            for ci, cell in enumerate(row.cells):
                master = cell_to_master.get((ri, ci), (ri, ci))
                is_merged = master != (ri, ci)
                cell_info = {
                    "col": ci,
                    "text": cell.text[:50] + ("..." if len(cell.text) > 50 else "")
                }
                if show_merged and is_merged:
                    cell_info["merged_with"] = f"[{master[0]},{master[1]}]"
                cells_info.append(cell_info)
            rows_info.append({"row": ri, "cells": cells_info})
        
        return json.dumps({
            "table_index": table_index,
            "rows": len(table.rows),
            "cols": len(table.columns),
            "unique_cells": len(unique_cells),
            "content": rows_info
        }, ensure_ascii=False, indent=2)
    else:
        result = []
        for ti, table in enumerate(tables):
            unique_cells, _ = _get_unique_cells(table)
            result.append({
                "table_index": ti,
                "rows": len(table.rows),
                "cols": len(table.columns),
                "unique_cells": len(unique_cells)
            })
        return json.dumps({"total_tables": len(tables), "tables": result}, ensure_ascii=False, indent=2)


def set_table_cell(filename, table_index, row_index, col_index, text, use_visual_index=False):
    """Set the text content of a specific table cell.
    
    If use_visual_index=True, col_index refers to the visual column position,
    skipping merged cells. This is useful for tables with complex merging.
    """
    _check_docx()
    doc = Document(filename)
    tables = doc.tables
    
    if table_index >= len(tables):
        return f"Error: Table index {table_index} out of range (total: {len(tables)})"
    
    table = tables[table_index]
    if row_index >= len(table.rows):
        return f"Error: Row index {row_index} out of range (total: {len(table.rows)})"
    
    row = table.rows[row_index]
    
    if use_visual_index:
        # Find the actual cell at the visual position
        unique_cells, cell_to_master = _get_unique_cells(table)
        visual_col = 0
        target_cell = None
        seen_in_row = set()
        
        for ci, cell in enumerate(row.cells):
            cell_id = id(cell)
            if cell_id not in seen_in_row:
                seen_in_row.add(cell_id)
                if visual_col == col_index:
                    target_cell = cell
                    break
                visual_col += 1
        
        if target_cell is None:
            return f"Error: Visual column {col_index} not found in row {row_index}"
        
        target_cell.text = str(text)
    else:
        if col_index >= len(row.cells):
            return f"Error: Column index {col_index} out of range (total: {len(row.cells)})"
        row.cells[col_index].text = str(text)
    
    doc.save(filename)
    return f"Set cell [{row_index},{col_index}] in table {table_index}"


def batch_set_table_cells(filename, table_index, cells_data, use_visual_index=False):
    """Set multiple table cells at once.
    
    cells_data: list of {"row": int, "col": int, "text": str}
    
    If use_visual_index=True, col refers to the visual column position,
    which is more intuitive for tables with merged cells.
    """
    _check_docx()
    doc = Document(filename)
    tables = doc.tables
    
    if table_index >= len(tables):
        return f"Error: Table index {table_index} out of range (total: {len(tables)})"
    
    table = tables[table_index]
    unique_cells, cell_to_master = _get_unique_cells(table)
    
    success_count = 0
    errors = []
    
    for item in cells_data:
        row_idx = item.get("row")
        col_idx = item.get("col")
        text = item.get("text", "")
        
        if row_idx is None or col_idx is None:
            errors.append(f"Missing row or col in item: {item}")
            continue
            
        if row_idx >= len(table.rows):
            errors.append(f"Row {row_idx} out of range")
            continue
        
        row = table.rows[row_idx]
        
        if use_visual_index:
            # Find the actual cell at the visual position
            visual_col = 0
            target_cell = None
            seen_in_row = set()
            
            for ci, cell in enumerate(row.cells):
                cell_id = id(cell)
                if cell_id not in seen_in_row:
                    seen_in_row.add(cell_id)
                    if visual_col == col_idx:
                        target_cell = cell
                        break
                    visual_col += 1
            
            if target_cell is None:
                errors.append(f"Visual column {col_idx} not found in row {row_idx}")
                continue
            
            target_cell.text = str(text)
        else:
            if col_idx >= len(row.cells):
                errors.append(f"Col {col_idx} out of range in row {row_idx}")
                continue
            row.cells[col_idx].text = str(text)
        
        success_count += 1
    
    doc.save(filename)
    
    if errors:
        return json.dumps({"success": success_count, "errors": errors}, ensure_ascii=False)
    return f"Set {success_count} cells in table {table_index}"


def add_table_row(filename, table_index, position="end", copy_style_from=None):
    """Add a new row to a table.
    
    position: "end" (default), "start", or an integer row index to insert before
    copy_style_from: row index to copy formatting from (optional)
    """
    _check_docx()
    doc = Document(filename)
    tables = doc.tables
    
    if table_index >= len(tables):
        return f"Error: Table index {table_index} out of range (total: {len(tables)})"
    
    table = tables[table_index]
    
    if position == "end":
        new_row = table.add_row()
    elif position == "start":
        new_row = table.add_row()
        tbl = table._tbl
        tr = new_row._tr
        tbl.remove(tr)
        tbl.insert(0, tr)
    else:
        try:
            insert_idx = int(position)
            if insert_idx >= len(table.rows):
                new_row = table.add_row()
            else:
                new_row = table.add_row()
                tbl = table._tbl
                tr = new_row._tr
                tbl.remove(tr)
                tbl.insert(insert_idx, tr)
        except ValueError:
            return f"Error: Invalid position: {position}"
    
    doc.save(filename)
    return f"Added row to table {table_index} at {position}"


def add_table_rows(filename, table_index, count, position="end"):
    """Add multiple rows to a table."""
    _check_docx()
    doc = Document(filename)
    tables = doc.tables
    
    if table_index >= len(tables):
        return f"Error: Table index {table_index} out of range (total: {len(tables)})"
    
    table = tables[table_index]
    
    for _ in range(count):
        if position == "end":
            table.add_row()
        elif position == "start":
            new_row = table.add_row()
            tbl = table._tbl
            tr = new_row._tr
            tbl.remove(tr)
            tbl.insert(0, tr)
        else:
            try:
                insert_idx = int(position)
                new_row = table.add_row()
                tbl = table._tbl
                tr = new_row._tr
                tbl.remove(tr)
                tbl.insert(min(insert_idx, len(table.rows) - 1), tr)
            except ValueError:
                return f"Error: Invalid position: {position}"
    
    doc.save(filename)
    return f"Added {count} rows to table {table_index}"


def delete_table_row(filename, table_index, row_index):
    """Delete a row from a table."""
    _check_docx()
    doc = Document(filename)
    tables = doc.tables
    
    if table_index >= len(tables):
        return f"Error: Table index {table_index} out of range (total: {len(tables)})"
    
    table = tables[table_index]
    
    if row_index >= len(table.rows):
        return f"Error: Row index {row_index} out of range (total: {len(table.rows)})"
    
    tr = table.rows[row_index]._tr
    table._tbl.remove(tr)
    
    doc.save(filename)
    return f"Deleted row {row_index} from table {table_index}"


def delete_table_rows(filename, table_index, row_indices):
    """Delete multiple rows from a table.
    
    row_indices: list of row indices to delete (will be processed in reverse order)
    """
    _check_docx()
    doc = Document(filename)
    tables = doc.tables
    
    if table_index >= len(tables):
        return f"Error: Table index {table_index} out of range (total: {len(tables)})"
    
    table = tables[table_index]
    
    # Sort in reverse order to avoid index shifting
    sorted_indices = sorted(row_indices, reverse=True)
    deleted = 0
    
    for row_idx in sorted_indices:
        if row_idx < len(table.rows):
            tr = table.rows[row_idx]._tr
            table._tbl.remove(tr)
            deleted += 1
    
    doc.save(filename)
    return f"Deleted {deleted} rows from table {table_index}"


def add_table_column(filename, table_index, position="end", width=None):
    """Add a new column to a table.
    
    position: "end" (default), "start", or an integer column index to insert before
    width: column width in inches (optional)
    """
    _check_docx()
    doc = Document(filename)
    tables = doc.tables
    
    if table_index >= len(tables):
        return f"Error: Table index {table_index} out of range (total: {len(tables)})"
    
    table = tables[table_index]
    
    # Determine insert position
    if position == "end":
        col_idx = len(table.columns)
    elif position == "start":
        col_idx = 0
    else:
        try:
            col_idx = int(position)
        except ValueError:
            return f"Error: Invalid position: {position}"
    
    # Add a cell to each row
    for row in table.rows:
        # Get the row's XML element
        tr = row._tr
        # Create a new cell
        tc = OxmlElement('w:tc')
        
        # Add width if specified
        if width:
            tcPr = OxmlElement('w:tcPr')
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(int(width * 1440)))  # Convert inches to twips
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
            tc.append(tcPr)
        
        # Add an empty paragraph to the cell
        p = OxmlElement('w:p')
        tc.append(p)
        
        # Insert at the correct position
        existing_cells = tr.findall(qn('w:tc'))
        if col_idx >= len(existing_cells):
            tr.append(tc)
        else:
            tr.insert(list(tr).index(existing_cells[col_idx]), tc)
    
    doc.save(filename)
    return f"Added column to table {table_index} at position {position}"


def delete_table_column(filename, table_index, col_index):
    """Delete a column from a table."""
    _check_docx()
    doc = Document(filename)
    tables = doc.tables
    
    if table_index >= len(tables):
        return f"Error: Table index {table_index} out of range (total: {len(tables)})"
    
    table = tables[table_index]
    
    if col_index >= len(table.columns):
        return f"Error: Column index {col_index} out of range (total: {len(table.columns)})"
    
    # Remove the cell at col_index from each row
    for row in table.rows:
        tr = row._tr
        cells = tr.findall(qn('w:tc'))
        if col_index < len(cells):
            tr.remove(cells[col_index])
    
    doc.save(filename)
    return f"Deleted column {col_index} from table {table_index}"


def add_hyperlink(filename, text, url, paragraph_index=None, color="0000FF", underline=True):
    """Add a hyperlink to the document.
    
    If paragraph_index is provided, appends to that paragraph.
    Otherwise, creates a new paragraph with the hyperlink.
    """
    _check_docx()
    from docx.shared import RGBColor
    
    doc = Document(filename)
    
    if paragraph_index is not None:
        paragraphs = list(doc.paragraphs)
        if paragraph_index < 0 or paragraph_index >= len(paragraphs):
            return f"Error: paragraph_index {paragraph_index} out of range"
        p = paragraphs[paragraph_index]
    else:
        p = doc.add_paragraph()
    
    # Create hyperlink element
    part = doc.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    # Create run with text
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Set color
    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color.lstrip('#'))
        rPr.append(c)
    
    # Set underline
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
    
    new_run.append(rPr)
    
    # Add text
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    
    hyperlink.append(new_run)
    p._p.append(hyperlink)
    
    doc.save(filename)
    return f"Added hyperlink: {text} -> {url}"


def add_hyperlink_to_text(filename, find_text, url, color="0000FF", underline=True):
    """Convert existing text to a hyperlink.
    
    Finds the first occurrence of find_text and converts it to a hyperlink.
    """
    _check_docx()
    
    doc = Document(filename)
    
    for para in doc.paragraphs:
        if find_text in para.text:
            # Find the run containing the text
            for run in para.runs:
                if find_text in run.text:
                    # Create hyperlink
                    part = doc.part
                    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
                    
                    hyperlink = OxmlElement('w:hyperlink')
                    hyperlink.set(qn('r:id'), r_id)
                    
                    # Create new run with the text
                    new_run = OxmlElement('w:r')
                    rPr = OxmlElement('w:rPr')
                    
                    if color:
                        c = OxmlElement('w:color')
                        c.set(qn('w:val'), color.lstrip('#'))
                        rPr.append(c)
                    
                    if underline:
                        u = OxmlElement('w:u')
                        u.set(qn('w:val'), 'single')
                        rPr.append(u)
                    
                    new_run.append(rPr)
                    
                    t = OxmlElement('w:t')
                    t.text = find_text
                    new_run.append(t)
                    
                    hyperlink.append(new_run)
                    
                    # Replace the run text and insert hyperlink
                    before_text = run.text.split(find_text)[0]
                    after_text = run.text.split(find_text, 1)[1] if find_text in run.text else ""
                    
                    run.text = before_text
                    run._r.addnext(hyperlink)
                    
                    if after_text:
                        after_run = para.add_run(after_text)
                        hyperlink.addnext(after_run._r)
                    
                    doc.save(filename)
                    return f"Converted '{find_text}' to hyperlink"
    
    return f"Error: Text '{find_text}' not found"


def get_hyperlinks(filename):
    """Get all hyperlinks in the document."""
    _check_docx()
    import json
    
    doc = Document(filename)
    hyperlinks = []
    
    # Get relationships
    rels = doc.part.rels
    
    for para_idx, para in enumerate(doc.paragraphs):
        for elem in para._p:
            if elem.tag.endswith('}hyperlink'):
                r_id = elem.get(qn('r:id'))
                if r_id and r_id in rels:
                    rel = rels[r_id]
                    text = ''.join(t.text for t in elem.iter() if t.tag.endswith('}t'))
                    hyperlinks.append({
                        "paragraph_index": para_idx,
                        "text": text,
                        "url": rel.target_ref
                    })
    
    return json.dumps(hyperlinks, ensure_ascii=False, indent=2)
