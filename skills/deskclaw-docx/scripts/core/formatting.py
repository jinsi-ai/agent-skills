"""Text and table formatting."""

from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.table import _Cell
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None


def _check_docx():
    if Document is None:
        raise RuntimeError("python-docx is required. Install with: pip install python-docx")


def _hex_to_rgb(hex_color):
    h = str(hex_color).lstrip("#")
    if len(h) == 6:
        return tuple(int(h[i : i + 2], 16) for i in (0, 2, 4))
    return (0, 0, 0)


def format_text(filename, paragraph_index, start_pos, end_pos, bold=None, italic=None, underline=None, color=None, font_size=None, font_name=None):
    _check_docx()
    doc = Document(filename)
    paragraphs = list(doc.paragraphs)
    if paragraph_index < 0 or paragraph_index >= len(paragraphs):
        return f"Error: paragraph_index {paragraph_index} out of range (0-{len(paragraphs)-1})"
    p = paragraphs[paragraph_index]
    full_text = p.text
    if start_pos < 0 or end_pos > len(full_text) or start_pos >= end_pos:
        return "Error: invalid start_pos/end_pos"
    p.clear()
    if start_pos > 0:
        p.add_run(full_text[:start_pos])
    run = p.add_run(full_text[start_pos:end_pos])
    if bold is not None:
        run.font.bold = bool(bold)
    if italic is not None:
        run.font.italic = bool(italic)
    if underline is not None:
        run.font.underline = bool(underline)
    if color is not None:
        from docx.shared import RGBColor
        run.font.color.rgb = RGBColor(*_hex_to_rgb(color.lstrip("#")))
    if font_size is not None:
        run.font.size = Pt(font_size / 2.0)
    if font_name is not None:
        run.font.name = font_name
    if end_pos < len(full_text):
        p.add_run(full_text[end_pos:])
    doc.save(filename)
    return "Formatted text"


def search_and_replace(filename, find_text, replace_text):
    _check_docx()
    doc = Document(filename)
    count = 0
    for p in doc.paragraphs:
        if find_text in p.text:
            p.text = p.text.replace(find_text, replace_text)
            count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if find_text in cell.text:
                    cell.text = cell.text.replace(find_text, replace_text)
                    count += 1
    doc.save(filename)
    return f"Replaced {count} occurrence(s)"


def delete_paragraph(filename, paragraph_index):
    _check_docx()
    doc = Document(filename)
    body = doc.element.body
    paras = [e for e in body if e.tag.endswith("}p")]
    if paragraph_index < 0 or paragraph_index >= len(paras):
        return f"Error: paragraph_index {paragraph_index} out of range"
    body.remove(paras[paragraph_index])
    doc.save(filename)
    return "Deleted paragraph"


def create_custom_style(filename, style_name, bold=None, italic=None, font_size=None, font_name=None, color=None, base_style=None):
    _check_docx()
    doc = Document(filename)
    styles = doc.styles
    base = base_style or "Normal"
    if style_name not in styles:
        style = styles.add_style(style_name, 1)
        style.base_style = styles[base]
    else:
        style = styles[style_name]
    font = style.font
    if bold is not None:
        font.bold = bool(bold)
    if italic is not None:
        font.italic = bool(italic)
    if font_size is not None:
        font.size = Pt(font_size / 2.0)
    if font_name is not None:
        font.name = font_name
    if color is not None:
        from docx.shared import RGBColor
        font.color.rgb = RGBColor(*_hex_to_rgb(color.lstrip("#")))
    doc.save(filename)
    return f"Created/updated style: {style_name}"


def _get_table(doc, table_index):
    tables = list(doc.tables)
    if table_index < 0 or table_index >= len(tables):
        return None
    return tables[table_index]


def format_table(filename, table_index, has_header_row=None, border_style=None, shading=None):
    _check_docx()
    doc = Document(filename)
    table = _get_table(doc, table_index)
    if table is None:
        return f"Error: table_index {table_index} out of range"
    if has_header_row is not None and has_header_row and len(table.rows) > 0:
        for cell in table.rows[0].cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.bold = True
    doc.save(filename)
    return "Formatted table"


def set_table_cell_shading(filename, table_index, row_index, col_index, fill_color, pattern="clear"):
    _check_docx()
    doc = Document(filename)
    table = _get_table(doc, table_index)
    if table is None:
        return f"Error: table_index {table_index} out of range"
    try:
        cell = table.rows[row_index].cells[col_index]
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill_color.lstrip("#"))
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)
    except IndexError:
        return "Error: row_index or col_index out of range"
    doc.save(filename)
    return "Set cell shading"


def apply_table_alternating_rows(filename, table_index, color1="FFFFFF", color2="F2F2F2"):
    _check_docx()
    doc = Document(filename)
    table = _get_table(doc, table_index)
    if table is None:
        return f"Error: table_index {table_index} out of range"
    for i, row in enumerate(table.rows):
        fill = color1 if i % 2 == 0 else color2
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), fill.lstrip("#"))
            shd.set(qn("w:val"), "clear")
            tcPr.append(shd)
    doc.save(filename)
    return "Applied alternating row colors"


def highlight_table_header(filename, table_index, header_color="4472C4", text_color="FFFFFF"):
    _check_docx()
    doc = Document(filename)
    table = _get_table(doc, table_index)
    if table is None or len(table.rows) == 0:
        return "Error: table not found or empty"
    from docx.shared import RGBColor
    header = table.rows[0]
    for cell in header.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), header_color.lstrip("#"))
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(*_hex_to_rgb(text_color))
    doc.save(filename)
    return "Highlighted table header"


def format_table_cell_text(filename, table_index, row_index, col_index, text_content=None, bold=None, italic=None, underline=None, color=None, font_size=None, font_name=None):
    _check_docx()
    doc = Document(filename)
    table = _get_table(doc, table_index)
    if table is None:
        return f"Error: table_index {table_index} out of range"
    try:
        cell = table.rows[row_index].cells[col_index]
        if text_content is not None:
            cell.text = text_content
        for p in cell.paragraphs:
            for run in p.runs:
                if bold is not None:
                    run.font.bold = bool(bold)
                if italic is not None:
                    run.font.italic = bool(italic)
                if underline is not None:
                    run.font.underline = bool(underline)
                if color is not None:
                    from docx.shared import RGBColor
                    run.font.color.rgb = RGBColor(*_hex_to_rgb(color.lstrip("#")))
                if font_size is not None:
                    run.font.size = Pt(font_size / 2.0)
                if font_name is not None:
                    run.font.name = font_name
    except IndexError:
        return "Error: row_index or col_index out of range"
    doc.save(filename)
    return "Formatted cell text"


def set_table_cell_padding(filename, table_index, row_index, col_index, top=None, bottom=None, left=None, right=None, unit="points"):
    _check_docx()
    doc = Document(filename)
    table = _get_table(doc, table_index)
    if table is None:
        return f"Error: table_index {table_index} out of range"
    try:
        cell = table.rows[row_index].cells[col_index]
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement("w:tcMar")
        for name, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
            if val is not None:
                el = OxmlElement(f"w:{name}")
                el.set(qn("w:w"), str(val))
                el.set(qn("w:type"), "dxa" if unit == "points" else "pct")
                tcMar.append(el)
        tcPr.append(tcMar)
    except IndexError:
        return "Error: row_index or col_index out of range"
    doc.save(filename)
    return "Set cell padding"


def set_table_cell_alignment(filename, table_index, row_index, col_index, horizontal="left", vertical="top"):
    _check_docx()
    doc = Document(filename)
    table = _get_table(doc, table_index)
    if table is None:
        return f"Error: table_index {table_index} out of range"
    try:
        cell = table.rows[row_index].cells[col_index]
        if horizontal in ("left", "center", "right"):
            cell.paragraphs[0].alignment = getattr(WD_ALIGN_PARAGRAPH, horizontal.upper(), WD_ALIGN_PARAGRAPH.LEFT)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        vAlign = OxmlElement("w:vAlign")
        vAlign.set(qn("w:val"), vertical)
        tcPr.append(vAlign)
    except IndexError:
        return "Error: row_index or col_index out of range"
    doc.save(filename)
    return "Set cell alignment"


def set_table_alignment_all(filename, table_index, horizontal="left", vertical="top"):
    _check_docx()
    doc = Document(filename)
    table = _get_table(doc, table_index)
    if table is None:
        return f"Error: table_index {table_index} out of range"
    for row in table.rows:
        for cell in row.cells:
            if cell.paragraphs:
                cell.paragraphs[0].alignment = getattr(WD_ALIGN_PARAGRAPH, horizontal.upper(), WD_ALIGN_PARAGRAPH.LEFT)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            vAlign = OxmlElement("w:vAlign")
            vAlign.set(qn("w:val"), vertical)
            tcPr.append(vAlign)
    doc.save(filename)
    return "Set table alignment"


def merge_table_cells(filename, table_index, start_row, start_col, end_row, end_col):
    _check_docx()
    doc = Document(filename)
    table = _get_table(doc, table_index)
    if table is None:
        return f"Error: table_index {table_index} out of range"
    try:
        cell = table.rows[start_row].cells[start_col]
        cell.merge(table.rows[end_row].cells[end_col])
    except (IndexError, Exception) as e:
        return f"Error: {e}"
    doc.save(filename)
    return "Merged cells"


def merge_table_cells_horizontal(filename, table_index, row_index, start_col, end_col):
    return merge_table_cells(filename, table_index, row_index, start_col, row_index, end_col)


def merge_table_cells_vertical(filename, table_index, col_index, start_row, end_row):
    return merge_table_cells(filename, table_index, start_row, col_index, end_row, col_index)


def set_table_column_width(filename, table_index, col_index, width, width_type="points"):
    _check_docx()
    doc = Document(filename)
    table = _get_table(doc, table_index)
    if table is None:
        return f"Error: table_index {table_index} out of range"
    try:
        col = table.columns[col_index]
        col.width = Pt(width) if width_type == "points" else Inches(width / 96.0)
    except (IndexError, Exception):
        table.rows[0].cells[col_index].width = Pt(width) if width_type == "points" else None
    doc.save(filename)
    return "Set column width"


def set_table_column_widths(filename, table_index, widths, width_type="points"):
    _check_docx()
    doc = Document(filename)
    table = _get_table(doc, table_index)
    if table is None:
        return f"Error: table_index {table_index} out of range"
    for i, w in enumerate(widths):
        if i < len(table.columns):
            try:
                table.columns[i].width = Pt(w) if width_type == "points" else None
            except Exception:
                pass
    doc.save(filename)
    return "Set column widths"


def set_table_width(filename, table_index, width, width_type="points"):
    _check_docx()
    doc = Document(filename)
    table = _get_table(doc, table_index)
    if table is None:
        return f"Error: table_index {table_index} out of range"
    table.autofit = False
    table.width = Pt(width) if width_type == "points" else None
    doc.save(filename)
    return "Set table width"


def auto_fit_table_columns(filename, table_index):
    _check_docx()
    doc = Document(filename)
    table = _get_table(doc, table_index)
    if table is None:
        return f"Error: table_index {table_index} out of range"
    table.autofit = True
    doc.save(filename)
    return "Auto-fit columns"


def set_paragraph_alignment(filename, paragraph_index, alignment):
    """Set paragraph alignment.
    
    alignment: 'left', 'center', 'right', 'justify'
    """
    _check_docx()
    doc = Document(filename)
    paragraphs = list(doc.paragraphs)
    
    if paragraph_index < 0 or paragraph_index >= len(paragraphs):
        return f"Error: paragraph_index {paragraph_index} out of range (0-{len(paragraphs)-1})"
    
    p = paragraphs[paragraph_index]
    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    p.alignment = align_map.get(alignment.lower(), WD_ALIGN_PARAGRAPH.LEFT)
    
    doc.save(filename)
    return f"Set paragraph {paragraph_index} alignment to {alignment}"


def set_paragraph_indent(filename, paragraph_index, left=None, right=None, first_line=None, hanging=None):
    """Set paragraph indentation.
    
    All values are in points (pt).
    left: left indent
    right: right indent
    first_line: first line indent (positive value)
    hanging: hanging indent (positive value, mutually exclusive with first_line)
    """
    _check_docx()
    doc = Document(filename)
    paragraphs = list(doc.paragraphs)
    
    if paragraph_index < 0 or paragraph_index >= len(paragraphs):
        return f"Error: paragraph_index {paragraph_index} out of range (0-{len(paragraphs)-1})"
    
    p = paragraphs[paragraph_index]
    pf = p.paragraph_format
    
    if left is not None:
        pf.left_indent = Pt(left)
    if right is not None:
        pf.right_indent = Pt(right)
    if first_line is not None:
        pf.first_line_indent = Pt(first_line)
    if hanging is not None:
        pf.first_line_indent = Pt(-hanging)
    
    doc.save(filename)
    return f"Set paragraph {paragraph_index} indentation"


def set_paragraph_spacing(filename, paragraph_index, before=None, after=None, line_spacing=None, line_spacing_rule=None):
    """Set paragraph spacing.
    
    before: space before paragraph (in points)
    after: space after paragraph (in points)
    line_spacing: line spacing value
    line_spacing_rule: 'single', 'one_point_five', 'double', 'exactly', 'at_least', 'multiple'
    """
    _check_docx()
    from docx.enum.text import WD_LINE_SPACING
    
    doc = Document(filename)
    paragraphs = list(doc.paragraphs)
    
    if paragraph_index < 0 or paragraph_index >= len(paragraphs):
        return f"Error: paragraph_index {paragraph_index} out of range (0-{len(paragraphs)-1})"
    
    p = paragraphs[paragraph_index]
    pf = p.paragraph_format
    
    if before is not None:
        pf.space_before = Pt(before)
    if after is not None:
        pf.space_after = Pt(after)
    
    if line_spacing_rule is not None:
        rule_map = {
            "single": WD_LINE_SPACING.SINGLE,
            "one_point_five": WD_LINE_SPACING.ONE_POINT_FIVE,
            "double": WD_LINE_SPACING.DOUBLE,
            "exactly": WD_LINE_SPACING.EXACTLY,
            "at_least": WD_LINE_SPACING.AT_LEAST,
            "multiple": WD_LINE_SPACING.MULTIPLE,
        }
        pf.line_spacing_rule = rule_map.get(line_spacing_rule.lower(), WD_LINE_SPACING.SINGLE)
    
    if line_spacing is not None:
        if line_spacing_rule in ("exactly", "at_least"):
            pf.line_spacing = Pt(line_spacing)
        else:
            pf.line_spacing = line_spacing
    
    doc.save(filename)
    return f"Set paragraph {paragraph_index} spacing"


def format_all_paragraphs(filename, alignment=None, left_indent=None, first_line_indent=None, space_before=None, space_after=None, line_spacing=None):
    """Apply formatting to all paragraphs in the document."""
    _check_docx()
    doc = Document(filename)
    
    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    
    for p in doc.paragraphs:
        pf = p.paragraph_format
        if alignment:
            p.alignment = align_map.get(alignment.lower(), WD_ALIGN_PARAGRAPH.LEFT)
        if left_indent is not None:
            pf.left_indent = Pt(left_indent)
        if first_line_indent is not None:
            pf.first_line_indent = Pt(first_line_indent)
        if space_before is not None:
            pf.space_before = Pt(space_before)
        if space_after is not None:
            pf.space_after = Pt(space_after)
        if line_spacing is not None:
            pf.line_spacing = line_spacing
    
    doc.save(filename)
    return f"Formatted {len(doc.paragraphs)} paragraphs"
