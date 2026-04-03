"""Tests for content.py: table operations, hyperlinks."""

import json
import sys
from pathlib import Path

import pytest

_SCRIPTS_DIR = Path(__file__).resolve().parent.parent
if str(_SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(_SCRIPTS_DIR))

from core import (
    add_table,
    add_paragraph,
    create_document,
    get_table_info,
    set_table_cell,
    batch_set_table_cells,
    add_table_row,
    add_table_rows,
    delete_table_row,
    delete_table_rows,
    add_table_column,
    delete_table_column,
    add_hyperlink,
    add_hyperlink_to_text,
    get_hyperlinks,
)
from docx import Document


# ---------------------------------------------------------------------------
# get_table_info
# ---------------------------------------------------------------------------

class TestGetTableInfo:
    def test_summary_all_tables(self, docx_with_table):
        result = json.loads(get_table_info(docx_with_table))
        assert result["total_tables"] == 1
        assert result["tables"][0]["rows"] == 3
        assert result["tables"][0]["cols"] == 3

    def test_single_table_detail(self, docx_with_table):
        result = json.loads(get_table_info(docx_with_table, table_index=0))
        assert result["rows"] == 3
        assert result["cols"] == 3
        assert result["content"][0]["cells"][0]["text"] == "A1"

    def test_table_index_out_of_range(self, docx_with_table):
        result = json.loads(get_table_info(docx_with_table, table_index=99))
        assert "error" in result


# ---------------------------------------------------------------------------
# set_table_cell / batch_set_table_cells
# ---------------------------------------------------------------------------

class TestSetTableCell:
    def test_set_single_cell(self, docx_with_table):
        result = set_table_cell(docx_with_table, table_index=0, row_index=1, col_index=1, text="UPDATED")
        assert "Set cell" in result
        doc = Document(docx_with_table)
        assert doc.tables[0].rows[1].cells[1].text == "UPDATED"

    def test_set_cell_out_of_range(self, docx_with_table):
        result = set_table_cell(docx_with_table, table_index=0, row_index=99, col_index=0, text="X")
        assert "Error" in result

    def test_batch_set_cells(self, docx_with_table):
        cells = [
            {"row": 0, "col": 0, "text": "X1"},
            {"row": 0, "col": 1, "text": "X2"},
            {"row": 2, "col": 2, "text": "X3"},
        ]
        result = batch_set_table_cells(docx_with_table, table_index=0, cells_data=cells)
        assert "3" in result
        doc = Document(docx_with_table)
        assert doc.tables[0].rows[0].cells[0].text == "X1"
        assert doc.tables[0].rows[2].cells[2].text == "X3"


# ---------------------------------------------------------------------------
# Table row operations
# ---------------------------------------------------------------------------

class TestTableRows:
    def test_add_row_end(self, docx_with_table):
        add_table_row(docx_with_table, table_index=0, position="end")
        doc = Document(docx_with_table)
        assert len(doc.tables[0].rows) == 4

    def test_add_row_start(self, docx_with_table):
        add_table_row(docx_with_table, table_index=0, position="start")
        doc = Document(docx_with_table)
        assert len(doc.tables[0].rows) == 4
        # Original first row data should now be at index 1
        assert doc.tables[0].rows[1].cells[0].text == "A1"

    def test_add_rows_bulk(self, docx_with_table):
        add_table_rows(docx_with_table, table_index=0, count=3, position="end")
        doc = Document(docx_with_table)
        assert len(doc.tables[0].rows) == 6

    def test_delete_row(self, docx_with_table):
        delete_table_row(docx_with_table, table_index=0, row_index=1)
        doc = Document(docx_with_table)
        assert len(doc.tables[0].rows) == 2
        assert doc.tables[0].rows[1].cells[0].text == "A3"

    def test_delete_rows_bulk(self, docx_with_table):
        delete_table_rows(docx_with_table, table_index=0, row_indices=[0, 2])
        doc = Document(docx_with_table)
        assert len(doc.tables[0].rows) == 1
        assert doc.tables[0].rows[0].cells[0].text == "A2"


# ---------------------------------------------------------------------------
# Table column operations
# ---------------------------------------------------------------------------

class TestTableColumns:
    def test_add_column_end(self, docx_with_table):
        add_table_column(docx_with_table, table_index=0, position="end")
        doc = Document(docx_with_table)
        from docx.oxml.ns import qn
        cells_in_first_row = doc.tables[0].rows[0]._tr.findall(qn("w:tc"))
        assert len(cells_in_first_row) == 4

    def test_delete_column(self, docx_with_table):
        delete_table_column(docx_with_table, table_index=0, col_index=1)
        doc = Document(docx_with_table)
        from docx.oxml.ns import qn
        cells_in_first_row = doc.tables[0].rows[0]._tr.findall(qn("w:tc"))
        assert len(cells_in_first_row) == 2
        def _cell_text(tc):
            return "".join(t.text or "" for t in tc.iter() if t.tag.endswith("}t"))
        assert _cell_text(cells_in_first_row[0]) == "A1"
        assert _cell_text(cells_in_first_row[1]) == "C1"


# ---------------------------------------------------------------------------
# Hyperlinks
# ---------------------------------------------------------------------------

class TestHyperlinks:
    def test_add_hyperlink_new_paragraph(self, tmp_docx):
        result = add_hyperlink(tmp_docx, "Visit OpenAI", "https://openai.com")
        assert "Added hyperlink" in result
        links = json.loads(get_hyperlinks(tmp_docx))
        assert len(links) >= 1
        assert links[0]["url"] == "https://openai.com"
        assert "Visit OpenAI" in links[0]["text"]

    def test_add_hyperlink_to_existing_paragraph(self, docx_with_paragraphs):
        result = add_hyperlink(
            docx_with_paragraphs, "click here", "https://example.com",
            paragraph_index=1,
        )
        assert "Added hyperlink" in result

    def test_hyperlink_to_text(self, docx_with_paragraphs):
        result = add_hyperlink_to_text(
            docx_with_paragraphs,
            find_text="Second paragraph",
            url="https://example.org",
        )
        assert "Converted" in result
        links = json.loads(get_hyperlinks(docx_with_paragraphs))
        assert any("https://example.org" in link["url"] for link in links)

    def test_get_hyperlinks_empty(self, tmp_docx):
        links = json.loads(get_hyperlinks(tmp_docx))
        assert links == []
