"""Tests for formatting.py: paragraph alignment, indent, spacing, bulk format."""

import sys
from pathlib import Path

import pytest

_SCRIPTS_DIR = Path(__file__).resolve().parent.parent
if str(_SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(_SCRIPTS_DIR))

from core import (
    set_paragraph_alignment,
    set_paragraph_indent,
    set_paragraph_spacing,
    format_all_paragraphs,
)
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


class TestParagraphAlignment:
    def test_set_center(self, docx_with_paragraphs):
        result = set_paragraph_alignment(docx_with_paragraphs, paragraph_index=1, alignment="center")
        assert "alignment" in result
        doc = Document(docx_with_paragraphs)
        assert doc.paragraphs[1].alignment == WD_ALIGN_PARAGRAPH.CENTER

    def test_set_justify(self, docx_with_paragraphs):
        set_paragraph_alignment(docx_with_paragraphs, paragraph_index=2, alignment="justify")
        doc = Document(docx_with_paragraphs)
        assert doc.paragraphs[2].alignment == WD_ALIGN_PARAGRAPH.JUSTIFY

    def test_out_of_range(self, docx_with_paragraphs):
        result = set_paragraph_alignment(docx_with_paragraphs, paragraph_index=999, alignment="left")
        assert "Error" in result


class TestParagraphIndent:
    def test_first_line_indent(self, docx_with_paragraphs):
        result = set_paragraph_indent(docx_with_paragraphs, paragraph_index=1, first_line=24)
        assert "indentation" in result
        doc = Document(docx_with_paragraphs)
        pf = doc.paragraphs[1].paragraph_format
        assert pf.first_line_indent is not None
        assert abs(pf.first_line_indent - Pt(24)) < Pt(1)


class TestParagraphSpacing:
    def test_set_spacing(self, docx_with_paragraphs):
        result = set_paragraph_spacing(
            docx_with_paragraphs, paragraph_index=1,
            before=12, after=6, line_spacing=1.5,
        )
        assert "spacing" in result
        doc = Document(docx_with_paragraphs)
        pf = doc.paragraphs[1].paragraph_format
        assert abs(pf.space_before - Pt(12)) < Pt(1)
        assert abs(pf.space_after - Pt(6)) < Pt(1)


class TestFormatAllParagraphs:
    def test_bulk_format(self, docx_with_paragraphs):
        result = format_all_paragraphs(
            docx_with_paragraphs,
            alignment="justify",
            first_line_indent=24,
            line_spacing=1.5,
        )
        assert "Formatted" in result

        doc = Document(docx_with_paragraphs)
        for p in doc.paragraphs:
            assert p.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
            pf = p.paragraph_format
            assert pf.first_line_indent is not None
            assert abs(pf.first_line_indent - Pt(24)) < Pt(1)
