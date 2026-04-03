"""Tests for document.py: page setup, header/footer, watermark, TOC, footnotes/endnotes."""

import json
import sys
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path

import pytest

_SCRIPTS_DIR = Path(__file__).resolve().parent.parent
if str(_SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(_SCRIPTS_DIR))

from core import (
    set_page_size,
    set_page_margins,
    set_page_orientation,
    get_page_settings,
    add_header,
    add_footer,
    add_page_number,
    set_different_first_page_header,
    add_header_image,
    add_watermark,
    add_table_of_contents,
    update_table_of_contents,
    add_footnote,
    add_endnote,
    add_heading,
    add_paragraph,
)
from docx import Document


# ---------------------------------------------------------------------------
# Page setup
# ---------------------------------------------------------------------------

class TestPageSize:
    def test_set_a4(self, tmp_docx):
        set_page_size(tmp_docx, paper="A4")
        settings = json.loads(get_page_settings(tmp_docx))
        assert abs(settings["page_width_inches"] - 8.27) < 0.1
        assert abs(settings["page_height_inches"] - 11.69) < 0.1

    def test_set_letter(self, tmp_docx):
        set_page_size(tmp_docx, paper="A4")
        result = set_page_size(tmp_docx, paper="Letter")
        assert "Set page size" in result
        settings = json.loads(get_page_settings(tmp_docx))
        assert abs(settings["page_width_inches"] - 8.5) < 0.1
        assert abs(settings["page_height_inches"] - 11.0) < 0.1

    def test_set_custom(self, tmp_docx):
        set_page_size(tmp_docx, paper="custom", width=10.0, height=14.0)
        settings = json.loads(get_page_settings(tmp_docx))
        assert abs(settings["page_width_inches"] - 10.0) < 0.1
        assert abs(settings["page_height_inches"] - 14.0) < 0.1

    def test_invalid_paper(self, tmp_docx):
        result = set_page_size(tmp_docx, paper="XXXL")
        assert "Error" in result

    def test_set_margins(self, tmp_docx):
        set_page_margins(tmp_docx, top=1.5, bottom=1.5, left=2.0, right=2.0)
        settings = json.loads(get_page_settings(tmp_docx))
        assert abs(settings["margins"]["top"] - 1.5) < 0.05
        assert abs(settings["margins"]["left"] - 2.0) < 0.05

    def test_set_landscape(self, tmp_docx):
        set_page_size(tmp_docx, paper="A4")
        set_page_orientation(tmp_docx, orientation="landscape")
        settings = json.loads(get_page_settings(tmp_docx))
        assert settings["orientation"] == "landscape"
        assert settings["page_width_inches"] > settings["page_height_inches"]


# ---------------------------------------------------------------------------
# Header / Footer
# ---------------------------------------------------------------------------

class TestHeaderFooter:
    def test_add_header_center(self, tmp_docx):
        add_header(tmp_docx, "Company Inc.", alignment="center", bold=True)
        doc = Document(tmp_docx)
        header_text = doc.sections[0].header.paragraphs[0].text
        assert "Company Inc." in header_text

    def test_add_footer_left(self, tmp_docx):
        add_footer(tmp_docx, "Confidential", alignment="left")
        doc = Document(tmp_docx)
        footer_text = doc.sections[0].footer.paragraphs[0].text
        assert "Confidential" in footer_text

    def test_add_page_number(self, tmp_docx):
        result = add_page_number(tmp_docx, position="footer", alignment="center", format_text="Page {page}")
        assert "Added page number" in result
        doc = Document(tmp_docx)
        footer = doc.sections[0].footer
        xml_str = footer._element.xml
        assert "PAGE" in xml_str or "fldChar" in xml_str

    def test_different_first_page_header(self, tmp_docx):
        result = set_different_first_page_header(tmp_docx, enabled=True)
        assert "enabled" in result
        doc = Document(tmp_docx)
        assert doc.sections[0].different_first_page_header_footer is True

    def test_header_and_footer_coexist(self, tmp_docx):
        add_header(tmp_docx, "Header Text")
        add_footer(tmp_docx, "Footer Text")
        doc = Document(tmp_docx)
        assert "Header Text" in doc.sections[0].header.paragraphs[0].text
        footer_texts = [p.text for p in doc.sections[0].footer.paragraphs]
        assert any("Footer Text" in t for t in footer_texts)


# ---------------------------------------------------------------------------
# Watermark
# ---------------------------------------------------------------------------

class TestWatermark:
    def test_add_text_watermark(self, tmp_docx):
        result = add_watermark(tmp_docx, "DRAFT", font_size=72, color="C0C0C0")
        assert "Added watermark" in result
        doc = Document(tmp_docx)
        header = doc.sections[0].header
        texts = [p.text for p in header.paragraphs]
        assert any("DRAFT" in t for t in texts)

    def test_add_image_watermark(self, tmp_docx, tmp_image):
        from core import add_watermark_image
        result = add_watermark_image(tmp_docx, tmp_image, width=2)
        assert "Added image watermark" in result


# ---------------------------------------------------------------------------
# Table of Contents
# ---------------------------------------------------------------------------

class TestTOC:
    def test_add_toc_at_start(self, tmp_docx):
        add_heading(tmp_docx, "Chapter 1", level=1)
        add_paragraph(tmp_docx, "Some content")
        result = add_table_of_contents(tmp_docx, title="Contents", max_level=3, position="start")
        assert "Added table of contents" in result

        doc = Document(tmp_docx)
        body_xml = doc.element.body.xml
        assert "TOC" in body_xml

    def test_update_toc(self, tmp_docx):
        add_heading(tmp_docx, "Chapter 1", level=1)
        add_table_of_contents(tmp_docx, title="TOC")
        result = update_table_of_contents(tmp_docx)
        assert "Marked TOC for update" in result


# ---------------------------------------------------------------------------
# Footnotes / Endnotes
# ---------------------------------------------------------------------------

class TestFootnotesEndnotes:
    def test_add_footnote(self, docx_with_paragraphs):
        result = add_footnote(docx_with_paragraphs, paragraph_index=1, footnote_text="Source: test data")
        assert "Added footnote" in result

        with zipfile.ZipFile(docx_with_paragraphs, "r") as z:
            assert "word/footnotes.xml" in z.namelist()
            with z.open("word/footnotes.xml") as f:
                content = f.read().decode("utf-8")
                assert "Source: test data" in content

    def test_add_endnote(self, docx_with_paragraphs):
        result = add_endnote(docx_with_paragraphs, paragraph_index=1, endnote_text="See appendix B")
        assert "Added endnote" in result

        with zipfile.ZipFile(docx_with_paragraphs, "r") as z:
            assert "word/endnotes.xml" in z.namelist()
            with z.open("word/endnotes.xml") as f:
                content = f.read().decode("utf-8")
                assert "See appendix B" in content

    def test_footnote_creates_xml_if_missing(self, tmp_docx):
        """Fresh doc has no footnotes.xml; add_footnote should create it."""
        add_paragraph(tmp_docx, "A paragraph")
        result = add_footnote(tmp_docx, paragraph_index=0, footnote_text="Auto-created")
        assert "Added footnote" in result
        with zipfile.ZipFile(tmp_docx, "r") as z:
            assert "word/footnotes.xml" in z.namelist()

    def test_footnote_invalid_paragraph(self, tmp_docx):
        result = add_footnote(tmp_docx, paragraph_index=999, footnote_text="Nope")
        assert "Error" in result or "out of range" in result
